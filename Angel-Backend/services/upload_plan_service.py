import os
import re
import json
from typing import Dict, Any, Optional
import PyPDF2
import docx
from docx import Document
import tempfile
from openai import AsyncOpenAI

client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))

async def process_uploaded_plan(file_path: str, file_extension: str) -> str:
    """
    Process uploaded business plan file and extract text content
    Supports: PDF, DOC, DOCX, TXT
    """
    try:
        if file_extension == '.pdf':
            return await extract_pdf_text(file_path)
        elif file_extension == '.docx':
            return await extract_docx_text(file_path)
        elif file_extension == '.doc':
            return await extract_doc_text(file_path)
        elif file_extension == '.txt':
            return await extract_txt_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
            
    except Exception as e:
        print(f"Error processing file: {e}")
        raise Exception(f"Failed to process file: {str(e)}")

async def extract_pdf_text(file_path: str) -> str:
    """Extract text from PDF file"""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
                
        return text.strip()
    except Exception as e:
        raise Exception(f"Error extracting PDF text: {str(e)}")

async def extract_docx_text(file_path: str) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
            
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
                
        return text.strip()
    except Exception as e:
        raise Exception(f"Error extracting DOCX text: {str(e)}")

async def extract_doc_text(file_path: str) -> str:
    """Extract text from DOC file (requires additional library)"""
    try:
        # For .doc files, we'd need python-docx2txt or similar
        # For now, return a message to convert to .docx
        raise Exception("DOC files are not supported. Please convert to DOCX format.")
    except Exception as e:
        raise Exception(f"Error extracting DOC text: {str(e)}")

async def extract_txt_text(file_path: str) -> str:
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read().strip()
    except Exception as e:
        raise Exception(f"Error extracting TXT text: {str(e)}")

async def extract_business_info_from_plan(content: str) -> Dict[str, Any]:
    """
    Extract structured business information from plan content using AI
    """
    try:
        prompt = f"""
        Analyze this business plan document and extract the following information in JSON format:

        {{
            "business_name": "Company name",
            "business_type": "Type of business (e.g., service, product, technology)",
            "industry": "Industry sector",
            "mission": "Mission statement",
            "vision": "Vision statement", 
            "tagline": "Company tagline or slogan",
            "target_market": "Primary target customers",
            "value_proposition": "What value the business provides",
            "revenue_model": "How the business makes money",
            "competitive_advantage": "What makes this business unique",
            "problem_solved": "What problem does the business solve",
            "solution": "How the business solves the problem",
            "market_size": "Market size or opportunity",
            "business_structure": "Legal structure (LLC, Corp, etc.)",
            "location": "Business location",
            "founding_year": "Year founded or planned founding",
            "team_size": "Current or planned team size",
            "funding_needs": "Funding requirements",
            "key_metrics": "Important business metrics",
            "goals": "Business goals and objectives"
        }}

        Business Plan Content:
        {content[:8000]}  # Limit content to avoid token limits

        Extract the information accurately. If information is not available, use null. Return only valid JSON.
        """

        response = await client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are an expert business analyst. Extract business information from documents and return structured JSON data."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=2000
        )

        extracted_text = response.choices[0].message.content.strip()
        
        # Clean up the response to ensure it's valid JSON
        if extracted_text.startswith('```json'):
            extracted_text = extracted_text[7:]
        if extracted_text.endswith('```'):
            extracted_text = extracted_text[:-3]
            
        business_info = json.loads(extracted_text)
        
        # Validate and clean the extracted data
        validated_info = {}
        for key, value in business_info.items():
            if value and isinstance(value, str) and len(value.strip()) > 0:
                validated_info[key] = value.strip()
            else:
                validated_info[key] = None
                
        return validated_info
        
    except json.JSONDecodeError as e:
        print(f"JSON decode error: {e}")
        return create_fallback_business_info(content)
    except Exception as e:
        print(f"Error extracting business info: {e}")
        return create_fallback_business_info(content)

def create_fallback_business_info(content: str) -> Dict[str, Any]:
    """Create basic business info when AI extraction fails"""
    return {
        "business_name": extract_basic_info(content, ["company", "business", "organization"]),
        "business_type": extract_basic_info(content, ["service", "product", "technology"]),
        "industry": extract_basic_info(content, ["industry", "sector", "market"]),
        "mission": extract_basic_info(content, ["mission", "purpose"]),
        "vision": extract_basic_info(content, ["vision", "goal"]),
        "tagline": extract_basic_info(content, ["tagline", "slogan"]),
        "target_market": extract_basic_info(content, ["customer", "client", "target"]),
        "value_proposition": extract_basic_info(content, ["value", "benefit", "advantage"]),
        "revenue_model": extract_basic_info(content, ["revenue", "income", "pricing"]),
        "competitive_advantage": extract_basic_info(content, ["competitive", "unique", "differentiation"]),
        "problem_solved": extract_basic_info(content, ["problem", "challenge", "issue"]),
        "solution": extract_basic_info(content, ["solution", "approach", "method"]),
        "market_size": None,
        "business_structure": extract_basic_info(content, ["LLC", "corporation", "partnership"]),
        "location": extract_basic_info(content, ["location", "address", "city"]),
        "founding_year": extract_basic_info(content, ["founded", "established", "started"]),
        "team_size": None,
        "funding_needs": extract_basic_info(content, ["funding", "investment", "capital"]),
        "key_metrics": None,
        "goals": extract_basic_info(content, ["goal", "objective", "target"])
    }

def extract_basic_info(content: str, keywords: list) -> Optional[str]:
    """Extract basic information using keyword matching"""
    content_lower = content.lower()
    
    for keyword in keywords:
        pattern = rf'{keyword}[:\s]*([^\n\r]{{10,100}})'
        match = re.search(pattern, content_lower)
        if match:
            return match.group(1).strip()
    
    return None

async def validate_business_plan_content(content: str) -> Dict[str, Any]:
    """
    Validate that the uploaded content is actually a business plan
    """
    try:
        validation_prompt = f"""
        Analyze this document and determine if it's a business plan. Return JSON with:
        {{
            "is_business_plan": true/false,
            "confidence": 0.0-1.0,
            "missing_sections": ["list of missing typical business plan sections"],
            "content_type": "description of what type of document this is",
            "recommendations": "suggestions for improvement"
        }}

        Document content:
        {content[:4000]}

        Typical business plan sections include: Executive Summary, Company Description, Market Analysis, Organization, Service/Product Line, Marketing, Financial Projections, etc.
        """

        response = await client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are an expert at analyzing business documents and plans."},
                {"role": "user", "content": validation_prompt}
            ],
            temperature=0.3,
            max_tokens=1000
        )

        validation_result = response.choices[0].message.content.strip()
        
        if validation_result.startswith('```json'):
            validation_result = validation_result[7:]
        if validation_result.endswith('```'):
            validation_result = validation_result[:-3]
            
        return json.loads(validation_result)
        
    except Exception as e:
        print(f"Error validating business plan: {e}")
        return {
            "is_business_plan": True,  # Default to accepting
            "confidence": 0.5,
            "missing_sections": [],
            "content_type": "Unknown document type",
            "recommendations": "Unable to validate document structure"
        }

async def analyze_plan_completeness(content: str, business_info: Dict[str, Any]) -> Dict[str, Any]:
    """
    Analyze the uploaded business plan and identify what information is missing
    Returns analysis with missing questions that need to be answered
    """
    try:
        # List of all 45 business plan questions to check against (9 sections restructured)
        business_plan_questions = [
            "What is your business name?",
            "What is your business tagline or mission statement?",
            "What problem does your business solve?",
            "What makes your business unique?",
            "Describe your core product or service in detail",
            "What are the key features and benefits of your product/service?",
            "What is your product development timeline?",
            "Who is your target market?",
            "What is the size of your target market?",
            "Who are your main competitors?",
            "How is your target market currently solving this problem?",
            "Where will your business be located?",
            "What are your space and facility requirements?",
            "What are your short-term operational needs?",
            "What suppliers or vendors will you need?",
            "What are your staffing needs?",
            "How will you price your product/service?",
            "What are your projected sales for the first year?",
            "What are your estimated startup costs?",
            "What are your estimated monthly operating expenses?",
            "When do you expect to break even?",
            "How much funding do you need to get started?",
            "What are your financial projections for years 1-3?",
            "How will you track and manage your finances?",
            "How will you reach your target customers?",
            "What is your sales process?",
            "What is your customer acquisition cost?",
            "What is your customer lifetime value?",
            "How will you build brand awareness and credibility?",
            "What partnerships or collaborations could help you reach more customers?",
            "What business structure will you use (LLC, Corporation, etc.)?",
            "What licenses and permits do you need?",
            "What insurance coverage do you need?",
            "How will you protect your intellectual property?",
            "What contracts and agreements will you need?",
            "How will you handle taxes and compliance?",
            "What data privacy and security measures will you implement?",
            "What are the key milestones you hope to achieve in the first year?",
            "What additional products or services could you offer in the future?",
            "How will you expand to new markets or customer segments?",
            "What partnerships or strategic alliances could accelerate your growth?",
            "What are the biggest risks and challenges your business might face?",
            "What contingency plans do you have for major risks or setbacks?",
            "What is your biggest concern or fear about launching this business?",
            "What additional considerations or final thoughts do you have about your business plan?"
        ]
        
        analysis_prompt = f"""
        Analyze this business plan document and compare it against the standard business plan questions.
        
        Business Plan Content (first 10000 characters):
        {content[:10000]}
        
        Extracted Business Information:
        {json.dumps(business_info, indent=2)}
        
        Standard Business Plan Questions:
        {json.dumps(business_plan_questions, indent=2)}
        
        Return a JSON object with:
        {{
            "summary": "Brief summary of what information is present in the plan",
            "completeness_score": 0.0-1.0,
            "found_information": {{
                "business_name": true/false,
                "mission_vision": true/false,
                "problem_solution": true/false,
                "target_market": true/false,
                "competitors": true/false,
                "financial_projections": true/false,
                "marketing_strategy": true/false,
                "operational_plan": true/false,
                "legal_structure": true/false,
                "risk_analysis": true/false
            }},
            "missing_questions": [
                {{
                    "question_number": 1-45,
                    "question_text": "Full question text",
                    "category": "Category (e.g., Financial, Marketing, Operations)",
                    "priority": "high/medium/low"
                }}
            ],
            "recommendations": "Overall recommendations for completing the plan"
        }}
        
        Only include questions in missing_questions if the information is truly missing or insufficient.
        Be specific about what's missing - don't just list all questions.
        """

        response = await client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are an expert business plan analyst. Analyze business plans and identify missing information accurately."},
                {"role": "user", "content": analysis_prompt}
            ],
            temperature=0.3,
            max_tokens=3000
        )

        analysis_text = response.choices[0].message.content.strip()
        
        # Clean up JSON response
        if analysis_text.startswith('```json'):
            analysis_text = analysis_text[7:]
        elif analysis_text.startswith('```'):
            analysis_text = analysis_text[3:]
        if analysis_text.endswith('```'):
            analysis_text = analysis_text[:-3]
            
        analysis_result = json.loads(analysis_text)
        
        return analysis_result
        
    except json.JSONDecodeError as e:
        print(f"JSON decode error in plan analysis: {e}")
        return {
            "summary": "Unable to fully analyze the plan. Please review manually.",
            "completeness_score": 0.5,
            "found_information": {},
            "missing_questions": [],
            "recommendations": "The plan was uploaded but could not be automatically analyzed. Please proceed with the business planning questions."
        }
    except Exception as e:
        print(f"Error analyzing plan completeness: {e}")
        return {
            "summary": "Error analyzing plan",
            "completeness_score": 0.5,
            "found_information": {},
            "missing_questions": [],
            "recommendations": "An error occurred during analysis. Please proceed with the business planning questions."
        }